from __future__ import annotations

import datetime as dt
from collections import Counter, defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook, load_workbook

from enacom_transcriptor.paths import LOGO_PATH, ensure_dirs


# =========================
# Excel helpers
# =========================

def ensure_excel_file(xlsx_path: str, sheets: dict[str, tuple[str, ...] | None]) -> None:
    """
    Asegura que exista el archivo XLSX y las hojas requeridas.

    - Si headers es None: crea la hoja sin encabezados.
    - Si headers no es None: escribe encabezados solo si la hoja está vacía.
    """
    ensure_dirs()
    p = Path(xlsx_path)
    p.parent.mkdir(parents=True, exist_ok=True)

    if p.exists():
        wb = load_workbook(str(p))
    else:
        wb = Workbook()

    for sheet_name, headers in sheets.items():
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            if wb.sheetnames == ["Sheet"] and wb["Sheet"]["A1"].value is None:
                ws = wb["Sheet"]
                ws.title = sheet_name
            else:
                ws = wb.create_sheet(sheet_name)

        if headers is not None and ws["A1"].value is None:
            ws.append(list(headers))

    wb.save(str(p))


def append_to_excel(
    xlsx_path: str,
    row: list,
    sheet_name: str = "Transcripción",
    headers: tuple[str, ...] = ("Inicio", "Fin", "Texto"),
) -> None:
    ensure_excel_file(xlsx_path, {sheet_name: headers})

    p = Path(xlsx_path)
    wb = load_workbook(str(p))
    ws = wb[sheet_name]
    ws.append(row)
    wb.save(str(p))


def write_infracciones_excel(
    xlsx_path: str,
    infracciones: list[dict] | None,
    sheet_name: str = "Infracciones",
    resumen_sheet: str = "Resumen_infracciones",
) -> None:
    """
    Escribe (reemplazando) la hoja de infracciones y una hoja resumen.
    """
    ensure_excel_file(xlsx_path, {sheet_name: ("Archivo", "Término", "Inicio", "Fin", "Texto")})

    p = Path(xlsx_path)
    wb = load_workbook(str(p))

    if sheet_name in wb.sheetnames:
        wb.remove(wb[sheet_name])
    ws = wb.create_sheet(sheet_name)
    ws.append(["Archivo", "Término", "Inicio", "Fin", "Texto"])

    infracciones = infracciones or []
    for inf in infracciones:
        ws.append([
            inf.get("archivo", ""),
            inf.get("termino", ""),
            inf.get("inicio", ""),
            inf.get("fin", ""),
            inf.get("texto", ""),
        ])

    if resumen_sheet in wb.sheetnames:
        wb.remove(wb[resumen_sheet])
    ws2 = wb.create_sheet(resumen_sheet)
    ws2.append(["Archivo", "Término", "Ocurrencias"])

    counter = defaultdict(Counter)
    for inf in infracciones:
        counter[inf.get("archivo", "")][inf.get("termino", "")] += 1

    for archivo, c in counter.items():
        for termino, n in c.most_common():
            ws2.append([archivo, termino, n])

    wb.save(str(p))


# =========================
# DOCX helpers
# =========================

def _template_path() -> Path:
    
    return LOGO_PATH.parent / "plantilla_enacom_limpia.docx"


def _load_doc() -> Document:
    tp = _template_path()
    if tp.exists():
        try:
            return Document(str(tp))
        except Exception:
            pass
    return Document()


def _add_logo_if_exists(doc: Document) -> None:
    try:
        if LOGO_PATH.exists():
            doc.add_picture(str(LOGO_PATH), width=Inches(1.35))
    except Exception:
        pass


def _add_heading_safe(doc: Document, text: str, level: int) -> None:
    """
    Agrega un heading sin romper si la plantilla no tiene estilos 'Heading X'.
    """
    style_name = f"Heading {level}"
    try:
        doc.add_paragraph(text, style=style_name)
    except Exception:
        p = doc.add_paragraph(text)
        if p.runs:
            p.runs[0].bold = True


def _kv_table(doc: Document, rows: Iterable[tuple[str, str]]) -> None:
    """
    Tabla simple de 2 columnas.
    """
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"

    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)


def generar_informe_word(
    titulo: str,
    docx_out_path: str,
    combinado: bool,
    meta: dict,
    txt_path: str | None = None,
    infracciones: list[dict] | None = None,
    files_info: list[dict] | None = None,
) -> str:
    """
    Genera informe DOCX (individual o combinado), robusto a plantillas sin estilos.
    """
    ensure_dirs()

    doc = _load_doc()
       
    _add_heading_safe(doc, "Transcriptor de audios a texto", level=1)
    doc.add_paragraph("Informe combinado (lote)" if combinado else "Informe individual")

    doc.add_paragraph("")

    _add_heading_safe(doc, "Datos del procesamiento", level=2)
    generado = meta.get("generado") or dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    _kv_table(doc, [
        ("Archivo/Lote", titulo),
        ("Generado", str(generado)),
        ("Modelo Whisper", str(meta.get("model_size", ""))),
        ("Idioma", str(meta.get("lang", "auto") or "auto")),
        ("Duración segmento (s)", str(meta.get("segment_duration", ""))),
        ("Cantidad de archivos", str(meta.get("total_files", ""))),
        ("Duración total", str(meta.get("total_duration_hhmmss", ""))),
    ])

    doc.add_paragraph("")

    # Transcripción
    _add_heading_safe(doc, "Transcripción", level=2)

    def _read_body_lines(p: str) -> list[str]:
        try:
            lines = Path(p).read_text(encoding="utf-8").splitlines()
        except Exception:
            return []
        body = [ln for ln in lines if ln.strip().startswith("[")]
        if not body:
            body = [ln for ln in lines if ln.strip()]
        return body

    if combinado:
        files_info = files_info or []
        if not files_info:
            doc.add_paragraph("(No hay archivos para mostrar.)")
        else:
            for j, info in enumerate(files_info):
                if j > 0:
                    doc.add_page_break()

                arch = info.get("archivo", f"archivo_{j}")
                dur = info.get("duracion_hhmmss", "")
                _add_heading_safe(doc, f"{arch} — Duración: {dur}", level=3)

                tp = info.get("txt_path")
                if not tp:
                    doc.add_paragraph("(Sin TXT asociado.)")
                    continue

                lines = _read_body_lines(tp)
                if not lines:
                    doc.add_paragraph("(Sin texto para mostrar.)")
                    continue

                for ln in lines:
                    doc.add_paragraph(ln)

    else:
        if not txt_path:
            doc.add_paragraph("(No se indicó txt_path.)")
        else:
            lines = _read_body_lines(txt_path)
            if not lines:
                doc.add_paragraph("(No se detectó texto en el archivo.)")
            else:
                for ln in lines:
                    doc.add_paragraph(ln)

    # Infracciones
    doc.add_page_break()
    _add_heading_safe(doc, "Infracciones detectadas", level=2)

    infracciones = infracciones or []
    if not infracciones:
        doc.add_paragraph("No se detectaron infracciones (según la configuración).")
    else:
        resumen = defaultdict(Counter)
        for inf in infracciones:
            resumen[inf.get("archivo", "")][inf.get("termino", "")] += 1

        doc.add_paragraph("Resumen de ocurrencias por archivo y término.")

        t = doc.add_table(rows=1, cols=3)
        h = t.rows[0].cells
        h[0].text = "Archivo"
        h[1].text = "Término"
        h[2].text = "Ocurrencias"

        for archivo, c in resumen.items():
            for termino, n in c.most_common():
                r = t.add_row().cells
                r[0].text = str(archivo)
                r[1].text = str(termino)
                r[2].text = str(n)

        doc.add_paragraph("")
        doc.add_paragraph("Detalle:")

        t2 = doc.add_table(rows=1, cols=5)
        h2 = t2.rows[0].cells
        h2[0].text = "Archivo"
        h2[1].text = "Término"
        h2[2].text = "Inicio"
        h2[3].text = "Fin"
        h2[4].text = "Texto"

        for inf in infracciones[:300]:
            r = t2.add_row().cells
            r[0].text = str(inf.get("archivo", ""))
            r[1].text = str(inf.get("termino", ""))
            r[2].text = str(inf.get("inicio", ""))
            r[3].text = str(inf.get("fin", ""))
            r[4].text = str(inf.get("texto", ""))

    out = Path(docx_out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)
